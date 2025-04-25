from fastapi import APIRouter, UploadFile, Form, HTTPException, File
import json
import os
import sys
import re
import io
import uuid
from typing import Optional
from pydantic import BaseModel
import pdfplumber
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import docx
import openai

# Create uploads directory if it doesn't exist
UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# Create a router instead of an app
router = APIRouter()

OPENAI_API_KEY = ""
openai.api_key = OPENAI_API_KEY  # Configure OpenAI client

class AnalysisResponse(BaseModel):
    success: bool
    resume_data: Optional[dict] = None
    analysis: Optional[dict] = None
    error: Optional[str] = None

# --- File Extraction Functions (Unchanged) ---
def extract_text_from_pdf(file_data):
    try:
        text = ""
        pdf_file = io.BytesIO(file_data)
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 10:
                    text += page_text + "\n"
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(cell for cell in row if cell)
                        if row_text.strip():
                            text += "\n" + row_text
        if len(text.strip()) < 50:
            images = convert_from_bytes(file_data)
            for img in images:
                text += pytesseract.image_to_string(img, lang='eng') + "\n"
        return text.strip()
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_data):
    try:
        docx_file = io.BytesIO(file_data)
        doc = docx.Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += "\n" + row_text
        return text.strip()
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text_from_image(file_data, filename):
    try:
        img = Image.open(io.BytesIO(file_data))
        text = pytesseract.image_to_string(img, lang='eng')
        return text.strip()
    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def extract_text_from_file(file_data, filename):
    print(f"Extracting text from file: {filename}")
    file_extension = os.path.splitext(filename)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_data)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_data)
    elif file_extension == '.txt':
        return file_data.decode('utf-8', errors='ignore')
    elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
        return extract_text_from_image(file_data, filename)
    else:
        return "Unsupported file format. Please upload PDF, DOCX, TXT, PNG, JPG, JPEG, TIFF, or BMP files."

# --- Text Preprocessing (Unchanged) ---
def preprocess_text(resume_text):
    resume_text = re.sub(r'\|\s*', ' | ', resume_text)
    resume_text = re.sub(r'\t+', ' ', resume_text)
    resume_text = re.sub(r'\s{2,}', ' ', resume_text)
    resume_text = re.sub(r'\|', '\n', resume_text)
    return resume_text.strip()

# --- Resume Parsing with OpenAI ---
def parse_resume(resume_text):
    print("Parsing resume text")
    resume_text = preprocess_text(resume_text)
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",  # Use a more robust model
            messages=[
                {
                    "role": "system",
                    "content": "You are a resume parsing expert. Your task is to extract structured information from resume text, including content from tables or unconventional layouts. Respond only with a JSON object containing the extracted data."
                },
                {
                    "role": "user",
                    "content": (
                        "Extract the following information from the resume text, formatted as JSON:\n"
                        "1. personal_details (name, email, phone, location, linkedin, website)\n"
                        "2. education (list of education items with degree, institution, dates, location, gpa if available)\n"
                        "3. skills (comprehensive list of all technical and soft skills mentioned - for each skill, include the name of the skill and any proficiency level mentioned)\n"
                        "4. experience (list of work experiences with company, title, dates, location, and description)\n\n"
                        f"Resume text:\n{resume_text}\n\n"
                        "Respond ONLY with the JSON object, no other text."
                    )
                }
            ],
            max_tokens=2000,
            temperature=0.0,  # Ensure deterministic output
            response_format={"type": "json_object"}
        )
        parsed_data = json.loads(response.choices[0].message.content)
        print("OpenAI parsing successful")
        
        # Process skills to ensure proper structure
        if "skills" in parsed_data:
            if parsed_data["skills"] and isinstance(parsed_data["skills"][0], str):
                skills_objects = []
                for skill in parsed_data["skills"]:
                    skill_obj = process_skill_text(skill)
                    skills_objects.append(skill_obj)
                parsed_data["skills"] = skills_objects
        else:
            parsed_data["skills"] = extract_skills_regex(resume_text)
        
        return parsed_data
    except Exception as e:
        print(f"OpenAI parsing error: {str(e)}")
        return {
            "personal_details": extract_personal_details(resume_text),
            "education": extract_education_regex(resume_text),
            "skills": extract_skills_regex(resume_text),
            "experience": extract_experience_regex(resume_text)
        }

# --- Personal Details Extraction (Unchanged) ---
def extract_personal_details(resume_text):
    personal_details = {
        "name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin": None,
        "website": None
    }
    
    resume_text = resume_text.lower()
    lines = resume_text.split('\n')
    
    email_pattern = r'(?:email|e-mail|contact)[:\s]*([a-za-z0-9._%+-]+@[a-za-z0-9.-]+\.[a-z]{2,})'
    for line in lines:
        email_match = re.search(email_pattern, line)
        if email_match:
            personal_details["email"] = email_match.group(1)
            break
    if not personal_details["email"]:
        email_pattern = r'\b[a-za-z0-9._%+-]+@[a-za-z0-9.-]+\.[a-z]{2,}\b'
        email_matches = re.findall(email_pattern, resume_text)
        if email_matches:
            personal_details["email"] = email_matches[0]
    
    phone_pattern = r'(?:phone|mobile|contact)[:\s]*((\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
    for line in lines:
        phone_match = re.search(phone_pattern, line)
        if phone_match:
            personal_details["phone"] = ''.join(phone_match.group(1).split())
            break
    if not personal_details["phone"]:
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_matches = re.findall(phone_pattern, resume_text)
        if phone_matches:
            personal_details["phone"] = ''.join(phone_matches[0]).replace(' ', '')
    
    name_pattern = r'(?:name|contact)[:\s]*([a-zA-Z\s]{2,40}(?![a-zA-Z\s]*university|college|company))'
    for line in lines:
        name_match = re.search(name_pattern, line, re.IGNORECASE)
        if name_match:
            personal_details["name"] = name_match.group(1).strip().title()
            break
    if not personal_details["name"]:
        for line in lines[:10]:
            words = line.strip().split()
            if len(words) >= 2 and all(w[0].isupper() for w in words[:2]) and not re.search(r'(university|college|inc|corp)', line, re.IGNORECASE):
                personal_details["name"] = ' '.join(words[:2]).title()
                break
    
    linkedin_pattern = r'(?:linkedin|linked in)[:\s]*(https?://(?:www\.)?linkedin\.com/in/[a-za-z0-9_-]+)'
    for line in lines:
        linkedin_match = re.search(linkedin_pattern, line, re.IGNORECASE)
        if linkedin_match:
            personal_details["linkedin"] = linkedin_match.group(1)
            break
    if not personal_details["linkedin"]:
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-za-z0-9_-]+'
        linkedin_matches = re.findall(linkedin_pattern, resume_text)
        if linkedin_matches:
            personal_details["linkedin"] = linkedin_matches[0]
    
    location_pattern = r'(?:address|location|city)[:\s]([a-zA-Z\s]+,\s[A-Z]{2}(?:\s*\d{5})?)'
    for line in lines:
        location_match = re.search(location_pattern, line, re.IGNORECASE)
        if location_match:
            personal_details["location"] = location_match.group(1).title()
            break
    if not personal_details["location"]:
        location_patterns = [
            r'\b[a-z][a-z\s]+,\s*[a-z]{2}\b',
            r'\b[a-z][a-z\s]+\s*[a-z]{2}\s*\d{5}\b'
        ]
        for pattern in location_patterns:
            location_matches = re.findall(pattern, resume_text)
            if location_matches:
                personal_details["location"] = location_matches[0].title()
                break
    
    website_pattern = r'(?:website|portfolio)[:\s](https?://(?:www\.)?[a-za-z0-9-]+\.[a-z]{2,}(?:/[^\s])?)'
    for line in lines:
        website_match = re.search(website_pattern, line, re.IGNORECASE)
        if website_match:
            personal_details["website"] = website_match.group(1)
            break
    if not personal_details["website"]:
        website_pattern = r'(?:https?://)?(?:www\.)?[a-za-z0-9-]+\.[a-z]{2,}(?:/[^\s]*)?'
        website_matches = re.findall(website_pattern, resume_text)
        for website in website_matches:
            if 'linkedin' not in website.lower() and 'github' not in website.lower():
                personal_details["website"] = website
                break
    
    return personal_details

# --- Section Finding (Unchanged) ---
def find_section(resume_text, section_titles):
    section_patterns = []
    for title in section_titles.split('|'):
        section_patterns.append(rf"(?:^|\n)(?:\s*)({title})(?:\s*(?::|\.|\n))")
    
    sections = []
    for pattern in section_patterns:
        matches = re.finditer(pattern, resume_text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            sections.append((match.start(), match.group(1)))
    
    sections.sort()
    
    if not sections:
        return ""
    
    section_texts = []
    for i, (start, title) in enumerate(sections):
        if i < len(sections) - 1:
            next_start = sections[i+1][0]
            section_text = resume_text[start:next_start]
        else:
            section_text = resume_text[start:]
        
        section_texts.append((title, section_text))
    
    for title, text in section_texts:
        for requested_title in section_titles.split('|'):
            if requested_title.lower() in title.lower():
                return text
    
    return ""

# --- Education Extraction (Unchanged) ---
def extract_education_regex(resume_text):
    education = []
    education_section = find_section(resume_text, "Education|Academic|Qualification|Educational Background")
    
    if not education_section:
        return education
    
    degree_patterns = [
        r'(B\.?Tech|Bachelor of Technology|M\.?Tech|Master of Technology|Ph\.?D|Doctor of Philosophy|MBA|Master of Business Administration|B\.?Sc|Bachelor of Science|M\.?Sc|Master of Science|B\.?E|Bachelor of Engineering|M\.?E|Master of Engineering|B\.?Com|Bachelor of Commerce|M\.?Com|Master of Commerce|B\.?A|Bachelor of Arts|M\.?A|Master of Arts)',
        r'(Bachelor|Master|Doctorate|Graduate|Post Graduate|Associate)'
    ]
    
    university_patterns = [
        r'University of [A-Za-z\s]+',
        r'[A-Z][a-zA-Z\s]+ University',
        r'[A-Z][a-zA-Z\s]+ College',
        r'[A-Z][a-zA-Z\s]+ Institute of [A-Za-z\s]+'
    ]
    
    paragraphs = re.split(r'\n\s*\n', education_section)
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if len(paragraph) > 10:
            has_degree = any(re.search(pattern, paragraph, re.IGNORECASE) for pattern in degree_patterns)
            has_university = any(re.search(pattern, paragraph, re.IGNORECASE) for pattern in university_patterns)
            has_year = re.search(r'(19|20)\d{2}', paragraph) is not None
            
            if (has_degree or has_university) and has_year:
                education.append(paragraph)
    
    if not education:
        bullet_points = re.findall(r'(?:^|\n)[•\-]\s(.?)(?=$|\n[•\-])', education_section, re.MULTILINE)
        for point in bullet_points:
            point = point.strip()
            if len(point) > 10:
                has_degree = any(re.search(pattern, point, re.IGNORECASE) for pattern in degree_patterns)
                has_university = any(re.search(pattern, point, re.IGNORECASE) for pattern in university_patterns)
                has_year = re.search(r'(19|20)\d{2}', point) is not None
                
                if (has_degree or has_university) and has_year:
                    education.append(point)
    
    return education

# --- Skills Extraction (Unchanged) ---
def extract_skills_regex(resume_text):
    skills = []
    processed_skills = set()
    
    table_skill_pattern = r'^([^\|\n]+?)(?:\s*\|\s*([^\|\n]?))?(?:\s\|\s*([^\|\n]*?))?$(?:\n)?'
    table_matches = re.finditer(table_skill_pattern, resume_text, re.MULTILINE)
    
    for match in table_matches:
        skill_name = match.group(1).strip()
        duration = match.group(2).strip() if match.group(2) else ""
        proficiency = match.group(3).strip() if match.group(3) else ""
        
        skill_obj = {
            "skill": skill_name,
            "proficiency_level": proficiency,
            "duration": duration
        }
        skill_key = skill_name.lower()
        if skill_key not in processed_skills:
            skills.append(skill_obj)
            processed_skills.add(skill_key)
    
    skills_section = find_section(resume_text, "Skills|Technical Skills|Skill Set|Technologies|Competencies|Expertise|Proficiencies") or resume_text
    
    bullet_skills = re.findall(r'(?:^|\n)[•\-]\s(.?)(?=$|\n[•\-])', skills_section, re.MULTILINE)
    for skill_text in bullet_skills:
        add_skill_to_list(skill_text, skills, processed_skills)
    
    comma_sections = re.findall(r'(?:^|\n|\:)([^\.!\?\n]{10,150}(?:,\s*[^,\.!\?\n]+){2,})(?:$|\n|\.|!|\?)', skills_section)
    for section in comma_sections:
        skill_items = [item.strip() for item in re.split(r',|\band\b', section) if len(item.strip()) > 0]
        for skill_text in skill_items:
            add_skill_to_list(skill_text, skills, processed_skills)
    
    if not skills:
        tech_terms = re.findall(r'\b([A-Za-z0-9#\+]+(?:\.[A-Za-z0-9]+)?)\b', skills_section)
        for term in tech_terms:
            if len(term) > 2 and term.lower() not in ['the', 'and', 'for', 'with', 'using']:
                add_skill_to_list(term, skills, processed_skills)
    
    return skills

def add_skill_to_list(skill_text, skills_list, processed_skills, category=None):
    skill_text = skill_text.strip()
    
    if not skill_text or len(skill_text) > 50 or re.search(r'\b(I|we|he|she|they|it)\b', skill_text, re.IGNORECASE):
        return
    
    skill_obj = process_skill_text(skill_text)
    
    if category and not re.search(fr'\b{re.escape(category)}\b', skill_obj["skill"], re.IGNORECASE):
        if not any(c in skill_obj["skill"].lower() for c in ["programming", "language", "framework", "database"]):
            skill_obj["category"] = category
    
    skill_key = skill_obj["skill"].lower()
    
    if skill_key and skill_key not in processed_skills:
        skills_list.append(skill_obj)
        processed_skills.add(skill_key)

def process_skill_text(skill_text):
    """Process a skill text to extract skill name, proficiency level, and duration."""
    # First, check if there's a parenthetical proficiency level
    paren_match = re.search(r'(.*?)\(([^)]+)\)', skill_text)
    if paren_match:
        base_skill = paren_match.group(1).strip()
        paren_content = paren_match.group(2).strip()
        
        # Check if parenthetical content looks like a proficiency level
        proficiency_indicators = ['basic', 'intermediate', 'advanced', 'beginner', 'expert', 
                                 'understanding', 'knowledge', 'familiar', 'proficient']
        
        if any(indicator in paren_content.lower() for indicator in proficiency_indicators):
            return {
                "skill": base_skill,
                "proficiency_level": paren_content,
                "duration": ""
            }
    
    # If no parenthetical proficiency, process normally
    skill_obj = {
        "skill": skill_text.strip(),
        "proficiency_level": "",
        "duration": ""
    }
    
    # Extended list of proficiency indicators
    proficiency_patterns = [
        r'\b(advanced|intermediate|beginner|expert|proficient|basic)\b',
        r'\b(working knowledge|basic understanding|familiar with)\b',
        r'\b(novice|experienced|professional level|advanced level|intermediate level|beginner level)\b'
    ]
    
    # Check each pattern for a match
    for pattern in proficiency_patterns:
        proficiency_match = re.search(pattern, skill_text, re.IGNORECASE)
        if proficiency_match:
            skill_obj["proficiency_level"] = proficiency_match.group(1)
            # Remove proficiency level from skill name
            skill_text_parts = re.split(pattern, skill_text, flags=re.IGNORECASE)
            skill_obj["skill"] = (skill_text_parts[0] + " " + "".join(skill_text_parts[2:])).strip()
            break
    
    # Check for duration (x years/months)
    duration_match = re.search(r'([0-9]+(?:\.[0-9]+)?\s*(?:years|year|yr|yrs|months|month|mo|mos))', skill_text, re.IGNORECASE)
    if duration_match:
        skill_obj["duration"] = duration_match.group(1)
        # Remove duration from skill name if not already modified by proficiency
        if skill_obj["skill"] == skill_text.strip() or skill_obj["skill"].endswith(skill_obj["proficiency_level"]):
            skill_obj["skill"] = re.sub(duration_match.group(0), "", skill_obj["skill"]).strip()
    
    # Clean up any trailing punctuation from the skill name
    skill_obj["skill"] = re.sub(r'[^\w\s]$', '', skill_obj["skill"]).strip()
    
    return skill_obj

# --- Experience Extraction (Unchanged) ---
def extract_experience_regex(resume_text):
    experiences = []
    experience_section = find_section(resume_text, "Experience|Work Experience|Employment|Professional Experience|Work History|Career History|Professional Background")
    
    if not experience_section:
        return experiences
    
    date_patterns = [
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]\.?\s\d{4}\s*(?:-|–|to)\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]\.?\s\d{4})',
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s*\d{4}\s*(?:-|–|to)\s*(?:January|February|March|April|May|June|July|August|September|October|November|December)\s*\d{4})',
        r'(\d{1,2}/\d{4}\s*(?:-|–|to)\s*\d{1,2}/\d{4})',
        r'(\d{4}\s*(?:-|–|to)\s*(?:Present|Current|Now|\d{4}))',
        r'(\d{1,2}/\d{4}\s*(?:-|–|to)\s*(?:Present|Current|Now|\d{1,2}/\d{4}))'
    ]
    
    job_entries = []
    
    for pattern in date_patterns:
        date_matches = list(re.finditer(pattern, experience_section, re.IGNORECASE))
        if date_matches:
            for i, match in enumerate(date_matches):
                start_pos = match.start()
                line_start = experience_section.rfind('\n', 0, start_pos)
                if line_start == -1:
                    line_start = 0
                else:
                    line_start += 1
                
                if i < len(date_matches) - 1:
                    next_date_pos = date_matches[i+1].start()
                    line_break_before_next = experience_section.rfind('\n', 0, next_date_pos)
                    if line_break_before_next > start_pos:
                        end_pos = line_break_before_next
                    else:
                        end_pos = next_date_pos
                else:
                    end_pos = len(experience_section)
                
                job_entry = experience_section[line_start:end_pos].strip()
                if job_entry and len(job_entry) > 30:
                    job_entries.append(job_entry)
    
    if not job_entries:
        paragraphs = re.split(r'\n\s*\n', experience_section)
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if len(paragraph) > 50:
                job_entries.append(paragraph)
    
    return job_entries

# --- Analysis with OpenAI ---
def analyze_with_openai(resume_text, job_description, job_role):
    print("Starting OpenAI analysis")
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are an ATS expert that analyzes resumes against job descriptions. Your task is to provide a detailed analysis including match percentage, matching keywords, missing keywords, and suggestions for improvement. Respond only with a JSON object containing the analysis."
                },
                {
                    "role": "user",
                    "content": (
                        "Analyze the resume against the job description and provide in JSON format:\n"
                        "1) overall_match_percentage (integer)\n"
                        "2) matching_keywords (list of strings)\n"
                        "3) missing_keywords (list of strings)\n"
                        "4) suggestions (list of strings)\n"
                        "5) strengths (list of strings)\n\n"
                        f"Job Role: {job_role}\n"
                        f"Job Description: {job_description}\n"
                        f"Resume: {resume_text}\n\n"
                        "Respond ONLY with the JSON object, no other text."
                    )
                }
            ],
            max_tokens=1500,
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        analysis_result = json.loads(response.choices[0].message.content)
        print("OpenAI analysis completed")
        return {
            "match_percentage": analysis_result.get("overall_match_percentage", 50),
            "matching_keywords": analysis_result.get("matching_keywords", []),
            "missing_keywords": analysis_result.get("missing_keywords", []),
            "suggestions": analysis_result.get("suggestions", []),
            "strengths": analysis_result.get("strengths", [])
        }
    except Exception as e:
        print(f"OpenAI error: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "match_percentage": 0,
            "matching_keywords": [],
            "missing_keywords": [],
            "suggestions": [],
            "strengths": []
        }

# --- Percentage Extraction (Unchanged) ---
def extract_percentage_from_analysis(analysis_text):
    percentage_pattern = r'(\d{1,3})%'
    match = re.search(percentage_pattern, analysis_text)
    if match:
        return int(match.group(1))
    return 50

# --- API Endpoints (Unchanged) ---
@router.get("/")
def read_root():
    return {
        "name": "Resume Analyzer API",
        "version": "1.0",
        "endpoints": {
            "/analyze": "POST - Submit resume for analysis",
            "/docs": "GET - API documentation"
        },
        "status": "running"
    }

@router.post("/openai/analyze", response_model=AnalysisResponse)
async def analyze_resume(
    resume_file: UploadFile = File(...),
    job_role: str = Form(""),
    job_description: str = Form("")
):
    allowed_extensions = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']
    file_extension = os.path.splitext(resume_file.filename)[1].lower()
    
    if file_extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file format. Allowed: {', '.join(allowed_extensions)}")
    
    try:
        print(f"Processing file: {resume_file.filename}")
        file_data = await resume_file.read()
        print(f"File data length: {len(file_data)}")
        
        file_id = str(uuid.uuid4())
        saved_path = os.path.join(UPLOAD_DIR, f"{file_id}{file_extension}")
        print(f"Saving to: {saved_path}")
        
        with open(saved_path, 'wb') as f:
            f.write(file_data)
        
        resume_text = extract_text_from_file(file_data, resume_file.filename)
        if resume_text.startswith("Error"):
            raise HTTPException(status_code=400, detail=resume_text)
        
        print(f"Extracted resume text length: {len(resume_text)}")
        resume_data = parse_resume(resume_text)
        
        response = {
            "success": True,
            "resume_data": resume_data,
        }
        
        if job_description.strip():
            print("Job description provided, performing analysis")
            analysis_result = analyze_with_openai(resume_text, job_description, job_role)
            response["analysis"] = analysis_result
        else:
            print("No job description provided, adding placeholder message")
            response["analysis"] = {"message": "Enter Job Role and Job Description for analysis"}
        
        return response
    except Exception as e:
        print(f"Error processing request: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")